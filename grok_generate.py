# grok_generate.py
import os
import time
import csv
import argparse
from pathlib import Path
from datetime import datetime
from typing import Optional, Tuple, List

from dotenv import load_dotenv

# OpenAI-compatible client pointing at xAI
try:
    from openai import OpenAI
except Exception:
    raise SystemExit(
        "OpenAI python client not found. Install with:\n  pip install openai python-dotenv python-docx"
    )

# DOCX writer
try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    raise SystemExit("python-docx not found. Install with: pip install python-docx")


# ------------------------- Config -------------------------
DEFAULT_BASE_DIR = Path(r"D:/Vald Data")  # Root: contains team folders
DEFAULT_MODEL = "grok-3-mini"
DEFAULT_RPM = 2  # 2 requests/minute (simple pacing)
AGE_GROUP_TEXT = "11–16 years old female"

LOG_CSV = "run_grok_log.csv"
FAIL_LIST = "failed_grok.txt"


# ------------------------- Utilities -------------------------
def now_iso() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


class RateLimiter:
    """Simple rate limiter: at most `rpm` requests per minute."""

    def __init__(self, rpm: int):
        self.min_interval = 60.0 / max(1, rpm)
        self.last_call: float = 0.0

    def wait(self):
        if self.last_call <= 0:
            return
        elapsed = time.time() - self.last_call
        sleep_for = self.min_interval - elapsed
        if sleep_for > 0:
            time.sleep(sleep_for)

    def stamp(self):
        self.last_call = time.time()


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def ensure_log_headers(log_path: Path) -> None:
    # Keep previous column order for compatibility
    if not log_path.exists():
        with log_path.open("w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(
                ["timestamp", "athlete", "folder", "status", "model", "outfile"]
            )


def append_log(log_path: Path, row: list) -> None:
    with log_path.open("a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(row)


def append_fail(fail_path: Path, athlete: str, reason: str) -> None:
    with fail_path.open("a", encoding="utf-8") as f:
        f.write(f"[{now_iso()}] {athlete} -> {reason}\n")


def build_prompt(athlete_name: str, analysis_md: str) -> str:
    """
    Compose the user prompt for Grok. Includes the analysis markdown as context.
    """
    user_prompt = f"""
Act as a physical coach and give me an 8 weeks training program for {athlete_name}
who is a female athlete. Her age group is {AGE_GROUP_TEXT}.
Don't make it too big. The training program should include at least:
- Program Overview
- Goals
- Weekly plan
- Progressions
- Monitoring and Safety Notes etc. 

Use the following analysis (generated from the athlete's test images) as context
to tailor the plan. If something is unclear, make sensible, coaching-appropriate assumptions. Don't create any table inside the training program, just plain text with headings. In the weekly plan, you are going to make Weeks 1-2, Weeks 3-4, Weeks 5-6, and Weeks 7-8 as a sub-heading. Do not put this sub-heading in bullet points .

--- BEGIN ATHLETE ANALYSIS (Markdown) ---
{analysis_md}
--- END ATHLETE ANALYSIS ---
""".strip()
    return user_prompt


def call_grok(client: OpenAI, model: str, prompt: str) -> str:
    """
    Send the prompt to Grok (OpenAI-compatible chat completions) and return the text.
    Includes light retry with exponential backoff.
    """
    backoff = 3
    last_err: Optional[Exception] = None

    for attempt in range(1, 4):
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an expert physical performance coach. "
                            "Return Markdown only. Keep the plan concise, structured, and actionable."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.7,
            )
            return resp.choices[0].message.content or ""
        except Exception as e:
            last_err = e
            time.sleep(backoff)
            backoff *= 2

    raise RuntimeError(f"Grok API call failed after retries: {last_err}")


# ------------------------- Markdown -> DOCX -------------------------
def markdown_to_docx(md_text: str, out_path: Path, title: Optional[str] = None) -> None:
    """
    Minimal Markdown-to-DOCX converter for headings and lists.
    Keeps things simple but produces a clean .docx.
    """
    doc = Document()

    # Optional title at the top
    if title:
        t = doc.add_heading(title, level=0)
        for run in t.runs:
            run.font.size = Pt(16)

    in_code_block = False
    for raw_line in md_text.splitlines():
        line = raw_line.rstrip("\n")

        # code fences
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            doc.add_paragraph()
            continue

        if in_code_block:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = "Consolas"
            continue

        # horizontal rule
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph().add_run("—" * 20)
            continue

        # headings: # .. ######
        if line.startswith("#"):
            hashes = len(line) - len(line.lstrip("#"))
            text = line[hashes:].strip()
            level = min(max(hashes, 1), 6)
            doc.add_heading(text or " ", level=level)
            continue

        # bullets
        if line.lstrip().startswith(("- ", "* ", "• ")):
            text = line.lstrip()[2:].strip() if len(line.lstrip()) >= 2 else ""
            doc.add_paragraph(text, style="List Bullet")
            continue

        # numbered list (simple 1. 2. 3. detection)
        stripped = line.lstrip()
        if any(
            stripped.startswith(prefix) for prefix in [f"{i}. " for i in range(1, 10)]
        ):
            text = stripped[stripped.find(".") + 1 :].strip()
            doc.add_paragraph(text, style="List Number")
            continue

        # blank line -> spacing
        if not line.strip():
            doc.add_paragraph()
            continue

        # regular paragraph
        doc.add_paragraph(line)

    doc.save(str(out_path))


# ------------------------- Team/Athlete discovery -------------------------
def list_team_dirs(base: Path) -> List[Path]:
    """Team folders are direct subfolders of base (sorted)."""
    teams = [p for p in base.iterdir() if p.is_dir()]
    return sorted(teams, key=lambda p: p.name.lower())


def list_athlete_dirs(team_dir: Path) -> List[Path]:
    """Athlete folders are direct subfolders of a team folder (sorted)."""
    athletes = [p for p in team_dir.iterdir() if p.is_dir()]
    return sorted(athletes, key=lambda p: p.name.lower())


def find_analysis_file(athlete_dir: Path, athlete_name: str) -> Optional[Path]:
    """
    Prefer the exact "{Athlete} Analysis.md". If not found, fall back to the first "* Analysis.md".
    """
    exact = athlete_dir / f"{athlete_name} Analysis.md"
    if exact.exists():
        return exact
    # fallback: any "* Analysis.md"
    candidates = sorted(athlete_dir.glob("* Analysis.md"))
    return candidates[0] if candidates else None


# ------------------------- Processing -------------------------
def process_athlete_folder(
    client: OpenAI,
    model: str,
    team_name: str,
    athlete_dir: Path,
    rl: RateLimiter,
    overwrite: bool = True,
) -> Tuple[bool, str]:
    """
    Returns (success, message). Writes the training program DOCX on success.
    """
    athlete_name = athlete_dir.name.strip()
    analysis_file = find_analysis_file(athlete_dir, athlete_name)
    out_file = athlete_dir / f"{athlete_name} 8 Weeks Training Program.docx"

    if not analysis_file or not analysis_file.exists():
        return False, f"Analysis file not found for athlete: {athlete_name}"

    try:
        analysis_md = read_text(analysis_file)
    except Exception as e:
        return False, f"Failed to read analysis: {e}"

    # Build prompt
    prompt = build_prompt(athlete_name, analysis_md)

    # Rate-limit
    rl.wait()
    try:
        result_md = call_grok(client, model, prompt)
    finally:
        rl.stamp()

    # Write output as DOCX (overwrite by default)
    if out_file.exists() and not overwrite:
        return True, f"Exists (overwrite=False): {out_file.name}"

    try:
        markdown_to_docx(
            result_md,
            out_file,
            title=f"{athlete_name} — 8 Weeks Training Program",
        )
    except Exception as e:
        return False, f"Failed to write DOCX: {e}"

    return True, str(out_file)


def main():
    load_dotenv()  # load .env (expects XAI_API_KEY)
    api_key = os.getenv("XAI_API_KEY")
    if not api_key:
        raise SystemExit(
            "Missing XAI_API_KEY in environment. Add it to your .env:\n  XAI_API_KEY=your_key_here"
        )

    parser = argparse.ArgumentParser(
        description=(
            "Generate 8-week training programs with Grok for each athlete folder under each team folder "
            "(DOCX output). Scans: Base → Team → Athlete."
        )
    )
    parser.add_argument(
        "--base-dir",
        default=str(DEFAULT_BASE_DIR),
        help="Root directory that contains TEAM folders (default: D:/Vald Data)",
    )
    parser.add_argument(
        "--model",
        default=DEFAULT_MODEL,
        help="xAI model name (default: grok-3-mini)",
    )
    parser.add_argument(
        "--rpm",
        type=int,
        default=DEFAULT_RPM,
        help="Requests per minute rate limit (default: 2)",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Scan and log what would be processed, but do not call the API or write files.",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        default=True,
        help="Overwrite existing training program files (default: True)",
    )
    args = parser.parse_args()

    base_dir = Path(args.base_dir)
    model = args.model
    rpm = max(1, args.rpm)
    dry_run = args.dry_run
    overwrite = args.overwrite

    if not base_dir.exists():
        raise SystemExit(f"Base directory not found: {base_dir}")

    # Set up OpenAI-compatible client for xAI
    client = OpenAI(api_key=api_key, base_url="https://api.x.ai/v1")

    # Logs
    log_csv = base_dir / LOG_CSV
    fail_list = base_dir / FAIL_LIST
    ensure_log_headers(log_csv)

    # Rate limiter
    rl = RateLimiter(rpm=rpm)

    # Discover teams & athletes
    teams = list_team_dirs(base_dir)
    total_teams = len(teams)
    total_athletes = sum(len(list_athlete_dirs(t)) for t in teams)

    print(
        f"[{now_iso()}] Scanning base: {base_dir}\n"
        f" - Teams found: {total_teams}\n"
        f" - Athlete folders (all teams): {total_athletes}\n"
    )

    # Process by team, then athlete
    team_idx = 0
    processed_count = 0

    for team in teams:
        team_idx += 1
        athletes = list_athlete_dirs(team)
        print(
            f"[{now_iso()}] Team {team_idx}/{total_teams}: {team.name} — Athletes: {len(athletes)}"
        )

        for a_idx, athlete_dir in enumerate(athletes, start=1):
            athlete = athlete_dir.name
            # Show which team & athlete we're processing
            print(f"  [{now_iso()}] ({a_idx}/{len(athletes)}) {team.name} → {athlete}")

            analysis_file = find_analysis_file(athlete_dir, athlete)
            out_file = athlete_dir / f"{athlete} 8 Weeks Training Program.docx"

            if not analysis_file or not analysis_file.exists():
                msg = f"SKIP - analysis file missing for {athlete}"
                print("     ", msg)
                append_log(
                    log_csv,
                    [
                        now_iso(),
                        athlete,
                        str(athlete_dir),
                        "missing-analysis",
                        model,
                        "",
                    ],
                )
                append_fail(fail_list, athlete, "analysis file missing")
                continue

            if dry_run:
                print("     DRY RUN - would call Grok and write:", out_file.name)
                append_log(
                    log_csv,
                    [
                        now_iso(),
                        athlete,
                        str(athlete_dir),
                        "dry-run",
                        model,
                        out_file.name,
                    ],
                )
                continue

            if out_file.exists() and not overwrite:
                print("     SKIP - output exists and overwrite=False:", out_file.name)
                append_log(
                    log_csv,
                    [
                        now_iso(),
                        athlete,
                        str(athlete_dir),
                        "skipped-exists",
                        model,
                        out_file.name,
                    ],
                )
                continue

            try:
                ok, msg = process_athlete_folder(
                    client, model, team.name, athlete_dir, rl, overwrite=overwrite
                )
                if ok:
                    print("     DONE ->", msg)
                    append_log(
                        log_csv,
                        [now_iso(), athlete, str(athlete_dir), "ok", model, msg],
                    )
                else:
                    print("     FAIL ->", msg)
                    append_log(
                        log_csv,
                        [now_iso(), athlete, str(athlete_dir), "fail", model, ""],
                    )
                    append_fail(fail_list, athlete, msg)
            except Exception as e:
                err = f"Unhandled error: {e}"
                print("     ERROR ->", err)
                append_log(
                    log_csv, [now_iso(), athlete, str(athlete_dir), "error", model, ""]
                )
                append_fail(fail_list, athlete, err)

            processed_count += 1

        # friendly spacing per team
        print()

    print(
        f"[{now_iso()}] All done. Teams: {total_teams} | Athletes visited: {total_athletes} | "
        f"Processed attempts: {processed_count} | Log: {log_csv.name} | Fail list: {fail_list.name}"
    )


if __name__ == "__main__":
    main()
